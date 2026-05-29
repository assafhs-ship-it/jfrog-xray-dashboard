from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# --- Styles helpers ---
def add_heading(doc, text, level=1, color=None, rtl=False):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if rtl:
        pPr = p._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'right')
        pPr.append(jc)
    return p

def add_body(doc, text, bold=False, rtl=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if rtl:
        pPr = p._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'right')
        pPr.append(jc)
    return p

def add_bullet(doc, text, rtl=False):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    if rtl:
        pPr = p._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'right')
        pPr.append(jc)
    return p

def add_separator(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '4CAF50')
    pBdr.append(bottom)
    pPr.append(pBdr)

JFROG_GREEN = (76, 175, 80)
DARK = (33, 33, 33)

# =====================================================================
# ENGLISH SECTION
# =====================================================================
add_heading(doc, 'Part 1 — Adoption Strategy', level=1, color=JFROG_GREEN)
add_heading(doc, 'JFrog Xray — Senior PM, Security Adoption', level=2)
add_body(doc, '')

# --- PERSONAS ---
add_heading(doc, '1. Personas', level=2, color=JFROG_GREEN)

add_heading(doc, 'Persona A: The DevOps Engineer ("The Builder")', level=3)
add_body(doc, 'Goals:', bold=True)
add_bullet(doc, 'Keep pipelines fast and automated')
add_bullet(doc, 'Reduce manual intervention and friction')
add_bullet(doc, 'Own the infrastructure, not the security policy')
add_body(doc, 'Pain Points:', bold=True)
add_bullet(doc, 'Xray adds latency to builds — every scan is friction')
add_bullet(doc, 'CVE noise: thousands of findings with no prioritization creates paralysis')
add_bullet(doc, 'Xray feels like someone else\'s responsibility ("I\'m not the security guy")')
add_bullet(doc, 'Setup is complex: connecting repos, writing policies, tuning thresholds — nobody showed them how')
add_body(doc, 'Why they\'re stuck: They installed Xray minimally to satisfy a checkbox. No policies were configured, '
         'and now either it\'s silently failing or generating alerts nobody reads.')

add_heading(doc, 'Persona B: The Security Engineer ("The Defender")', level=3)
add_body(doc, 'Goals:', bold=True)
add_bullet(doc, 'Get full visibility into what\'s running in production')
add_bullet(doc, 'Reduce real organizational risk and prove compliance to auditors')
add_bullet(doc, 'Enforce security standards across the SDLC')
add_body(doc, 'Pain Points:', bold=True)
add_bullet(doc, 'Doesn\'t own Artifactory — depends on DevOps to configure it correctly')
add_bullet(doc, 'No enforcement power without engineering buy-in')
add_bullet(doc, 'License violations are an afterthought until legal asks')
add_bullet(doc, 'Cannot report to CISO on actual security posture because scans aren\'t running')
add_body(doc, 'Why they\'re stuck: They want Xray to be meaningful but don\'t control the toolchain. '
         'Getting DevOps to act requires political capital they often don\'t have.')

add_heading(doc, 'Persona C: The Development Team Lead ("The Accountable One")', level=3)
add_body(doc, 'Goals:', bold=True)
add_bullet(doc, 'Ship features fast and keep the team unblocked')
add_bullet(doc, 'Avoid production security incidents that cause fires')
add_body(doc, 'Pain Points:', bold=True)
add_bullet(doc, 'Security feels like a blocker, not an enabler')
add_bullet(doc, 'No clarity on "fix now" vs. "can wait" — everything looks urgent')
add_bullet(doc, 'Remediations are not mapped to their sprint or backlog workflow')
add_bullet(doc, 'Nobody showed them what "done" looks like with Xray')
add_body(doc, 'Why they\'re stuck: ROI is invisible. The license was purchased top-down but no one was '
         'accountable for activation. They\'ve never seen a clear value moment.')

add_heading(doc, 'Bonus Persona: The CISO / VP Engineering (Economic Buyer)', level=3)
add_bullet(doc, 'Paid for the license, cannot measure what they got')
add_bullet(doc, 'Renewal is at risk if the platform never proved its value')

add_body(doc, '')

# --- JOURNEY MAP ---
add_heading(doc, '2. Journey Map', level=2, color=JFROG_GREEN)
add_body(doc, 'Key insight: "Stuck" is not one state — it is four distinct stages of being stuck. '
         'The job is to identify where each customer is and apply the right intervention.')

add_body(doc, 'Adoption Stage Model:', bold=True)
add_bullet(doc, 'Stage 0 — Installed, Dormant: Xray is connected to Artifactory but zero policies configured')
add_bullet(doc, 'Stage 1 — Configured, Ignored: Policies exist, alerts are firing, nobody is acting on them')
add_bullet(doc, 'Stage 2 — Active, Shallow: Some repos scanned, CVEs acknowledged but not blocking the pipeline')
add_bullet(doc, 'Stage 3 — Embedded: Xray blocks bad artifacts in CI/CD; teams actively remediate issues')
add_body(doc, 'Most stuck customers are at Stage 0 or Stage 1.')

add_body(doc, '"Think Like a Hacker" — Proactive Intervention Tactics:', bold=True)
add_bullet(doc, 'Trigger-based proactive outreach: When telemetry shows zero policies after 30 days, '
           'CS automatically gets an alert to reach out — don\'t wait for the customer to ask')
add_bullet(doc, 'Value moment manufacturing: Run a scan on the customer\'s top 5 repos and send a '
           '"here\'s what we found" report before they ask. Use real data to create urgency, not a sales pitch')
add_bullet(doc, 'CVE news-jacking: When a major CVE drops (e.g., Log4Shell, XZ Utils), reach out to dormant '
           'customers: "Had Xray been configured, you would have caught this before it hit production"')
add_bullet(doc, 'Internal champion activation: Identify the one security-conscious engineer inside the account. '
           'Arm them with a business case and a 1-hour setup guide to become the internal hero')
add_bullet(doc, 'Benchmarking / peer pressure: Show anonymized data — "Customers like you have X policies and '
           'Y repos scanned." Nobody wants to be below the median')
add_bullet(doc, 'Health score nudges: Define a Xray Adoption Health Score (0–100) and send automated nudges '
           'when a customer is close to the next milestone. Progress visibility creates momentum')

add_body(doc, '')

# --- 30-DAY PLAN ---
add_heading(doc, '3. 30-Day Plan (As New PM in This Role)', level=2, color=JFROG_GREEN)

add_heading(doc, 'Week 1 — Understand the Landscape', level=3)
add_bullet(doc, 'Pull the stuck customer list; segment by adoption stage (0–3)')
add_bullet(doc, 'Talk to 3 CS reps: what blockers do you hear most? What\'s the current playbook (if any)?')
add_bullet(doc, 'Identify which customer segment has the highest volume AND highest churn risk')

add_heading(doc, 'Week 2 — Diagnose the Real Blockers', level=3)
add_bullet(doc, 'Talk directly to 5 stuck customers across 2–3 segments')
add_bullet(doc, 'Goal: find the top 2–3 friction points that appear across accounts '
           '(e.g., "policy setup too complex", "CVE overload, no prioritization", "no DevOps buy-in")')
add_bullet(doc, 'Identify 1–2 champion contacts inside stuck accounts who would co-design a solution')

add_heading(doc, 'Week 3 — Run One Targeted Experiment', level=3)
add_bullet(doc, 'Define the minimum intervention for the most common stuck pattern')
add_bullet(doc, 'Run it with 3 customers: e.g., a 45-min "Xray Activation Sprint" call with a pre-built '
           'policy template and guided setup')
add_bullet(doc, 'Instrument it: measure time-to-first-policy-configured, time-to-first-CVE-resolved')

add_heading(doc, 'Week 4 — Socialize and Define the Playbook', level=3)
add_bullet(doc, 'Report early results to CS leadership')
add_bullet(doc, 'Propose 1 structured CS playbook: what triggers outreach, what the call looks like, '
           'what success looks like')
add_bullet(doc, 'Define the adoption health score that will power the CS dashboard')

add_body(doc, 'Core framing: The biggest risk is not that we don\'t know what to build — it\'s that CS has no '
         'repeatable motion. The first 30 days is about generating enough signal to give CS a playbook, not '
         'building a product from scratch.', bold=False)

add_body(doc, '')
add_separator(doc)
add_body(doc, '')

# =====================================================================
# HEBREW SECTION
# =====================================================================
add_heading(doc, 'חלק 1 — אסטרטגיית אימוץ', level=1, color=JFROG_GREEN, rtl=True)
add_heading(doc, 'JFrog Xray — מנהל מוצר בכיר, אימוץ אבטחה', level=2, rtl=True)
add_body(doc, '', rtl=True)

# --- פרסונות ---
add_heading(doc, '1. פרסונות', level=2, color=JFROG_GREEN, rtl=True)

add_heading(doc, 'פרסונה א\': מהנדס DevOps ("הבונה")', level=3, rtl=True)
add_body(doc, ':מטרות', bold=True, rtl=True)
add_bullet(doc, 'לשמור על ה-pipelines מהירים ואוטומטיים', rtl=True)
add_bullet(doc, 'לצמצם התערבות ידנית וחיכוך בתהליכים', rtl=True)
add_bullet(doc, 'להיות אחראי על התשתית, לא על מדיניות האבטחה', rtl=True)
add_body(doc, ':נקודות כאב', bold=True, rtl=True)
add_bullet(doc, 'Xray מוסיף זמן ריצה ל-builds — כל סריקה היא חיכוך', rtl=True)
add_bullet(doc, 'רעש CVE: אלפי ממצאים ללא תעדוף יוצרים שיתוק', rtl=True)
add_bullet(doc, 'Xray מרגיש כאחריות של מישהו אחר ("אני לא איש האבטחה")', rtl=True)
add_bullet(doc, 'ההגדרה מורכבת: חיבור repos, כתיבת מדיניות, כיוון סף התראות — אף אחד לא הסביר להם איך', rtl=True)
add_body(doc, 'מדוע הם תקועים: התקינו את Xray באופן מינימלי כדי לסמן ✓. לא הוגדרה מדיניות, ועכשיו המוצר '
         'או שלא עובד בשקט או שמייצר התראות שאף אחד לא קורא.', rtl=True)

add_heading(doc, 'פרסונה ב\': מהנדס אבטחה ("המגן")', level=3, rtl=True)
add_body(doc, ':מטרות', bold=True, rtl=True)
add_bullet(doc, 'להשיג ראות מלאה על מה שרץ בייצור', rtl=True)
add_bullet(doc, 'לצמצם סיכון ארגוני אמיתי ולהוכיח עמידה ברגולציה', rtl=True)
add_bullet(doc, 'לאכוף סטנדרטים של אבטחה לאורך מחזור חיי הפיתוח', rtl=True)
add_body(doc, ':נקודות כאב', bold=True, rtl=True)
add_bullet(doc, 'לא שולט ב-Artifactory — תלוי ב-DevOps להגדיר נכון', rtl=True)
add_bullet(doc, 'אין כוח אכיפה ללא שיתוף פעולה של ההנדסה', rtl=True)
add_bullet(doc, 'הפרות רישיון הן עניין שולי עד שהמחלקה המשפטית שואלת', rtl=True)
add_bullet(doc, 'לא יכול לדווח ל-CISO על מצב האבטחה כי הסריקות לא רצות', rtl=True)
add_body(doc, 'מדוע הם תקועים: רוצים ש-Xray יהיה משמעותי אבל לא שולטים בכלים. '
         'לגרום ל-DevOps לפעול דורש הון פוליטי שלרוב אין להם.', rtl=True)

add_heading(doc, 'פרסונה ג\': ראש צוות פיתוח ("האחראי")', level=3, rtl=True)
add_body(doc, ':מטרות', bold=True, rtl=True)
add_bullet(doc, 'לשלוח פיצ\'רים מהר ולשמור על הצוות לא חסום', rtl=True)
add_bullet(doc, 'להימנע מאירועי אבטחה בייצור שגורמים לשריפות', rtl=True)
add_body(doc, ':נקודות כאב', bold=True, rtl=True)
add_bullet(doc, 'אבטחה מרגישה כחסם, לא כמאפשר', rtl=True)
add_bullet(doc, 'אין בהירות על "לתקן עכשיו" לעומת "יכול לחכות" — הכל נראה דחוף', rtl=True)
add_bullet(doc, 'תיקונים לא ממופים ל-sprint או ל-backlog של הצוות', rtl=True)
add_bullet(doc, 'אף אחד לא הראה להם איך נראה "סיים" עם Xray', rtl=True)
add_body(doc, 'מדוע הם תקועים: ה-ROI בלתי נראה. הרישיון נרכש מלמעלה למטה אבל '
         'אף אחד לא היה אחראי על ההפעלה. הם מעולם לא חוו רגע ערך ברור.', rtl=True)

add_heading(doc, 'פרסונה בונוס: CISO / סמנכ"ל הנדסה (הקונה הכלכלי)', level=3, rtl=True)
add_bullet(doc, 'שילם עבור הרישיון, לא יכול למדוד מה קיבל', rtl=True)
add_bullet(doc, 'חידוש החוזה בסכנה אם הפלטפורמה לא הוכיחה ערך', rtl=True)

add_body(doc, '', rtl=True)

# --- מפת מסע ---
add_heading(doc, '2. מפת מסע לקוח', level=2, color=JFROG_GREEN, rtl=True)
add_body(doc, 'תובנה מרכזית: "תקוע" אינו מצב אחד — זהו ארבעה שלבים שונים של היתקעות. '
         'המשימה היא לזהות איפה כל לקוח נמצא ולהפעיל את ההתערבות הנכונה.', rtl=True)

add_body(doc, ':מודל שלבי האימוץ', bold=True, rtl=True)
add_bullet(doc, 'שלב 0 — מותקן, רדום: Xray מחובר ל-Artifactory אבל אפס מדיניות מוגדרת', rtl=True)
add_bullet(doc, 'שלב 1 — מוגדר, מתעלמים: מדיניות קיימת, התראות יורות, אף אחד לא פועל', rtl=True)
add_bullet(doc, 'שלב 2 — פעיל, רדוד: כמה repos נסרקים, CVEs מוכרים אבל לא חוסמים את ה-pipeline', rtl=True)
add_bullet(doc, 'שלב 3 — מוטמע: Xray חוסם artifacts בעייתיים ב-CI/CD; צוותים מתקנים בפועל', rtl=True)
add_body(doc, 'רוב הלקוחות התקועים נמצאים בשלב 0 או 1.', rtl=True)

add_body(doc, ':טקטיקות התערבות יצירתיות ("לחשוב כמו האקר")', bold=True, rtl=True)
add_bullet(doc, 'פנייה פרואקטיבית מבוססת טריגר: כאשר הטלמטריה מראה אפס מדיניות אחרי 30 יום, '
           'CS מקבל אוטומטית התראה לפנות — לא לחכות שהלקוח יבקש עזרה', rtl=True)
add_bullet(doc, 'יצירת רגע ערך: להריץ סריקה על 5 ה-repos המרכזיים של הלקוח ולשלוח '
           'דוח "הנה מה שמצאנו" לפני שביקשו. להשתמש בנתונים אמיתיים כדי ליצור דחיפות', rtl=True)
add_bullet(doc, 'ניצול חדשות CVE: כאשר CVE גדול מתפרסם (Log4Shell, XZ Utils), לפנות ללקוחות רדומים: '
           '"אם Xray היה מוגדר, הייתם תופסים את זה לפני שהגיע לייצור"', rtl=True)
add_bullet(doc, 'הפעלת אלוף פנימי: לזהות את המהנדס הבודד שאכפת לו מאבטחה בתוך הלקוח. '
           'לצייד אותו עם תיק עסקי ומדריך הגדרה של שעה כדי שיהפוך לגיבור פנימי', rtl=True)
add_bullet(doc, 'השוואה עמיתים / לחץ חברתי: להציג נתונים אנונימיים — '
           '"לקוחות כמוכם מוגדרים עם X מדיניות ו-Y repos נסרקים." אף אחד לא רוצה להיות מתחת לחציון', rtl=True)
add_bullet(doc, 'תזכורות מבוססות ציון: להגדיר Xray Adoption Health Score (0–100) ולשלוח '
           'התראות אוטומטיות כשלקוח קרוב למדרגה הבאה. ראות בנושא התקדמות יוצרת מומנטום', rtl=True)

add_body(doc, '', rtl=True)

# --- תוכנית 30 יום ---
add_heading(doc, '3. תוכנית 30 יום (כמנהל מוצר חדש בתפקיד)', level=2, color=JFROG_GREEN, rtl=True)

add_heading(doc, 'שבוע 1 — הבנת הנוף', level=3, rtl=True)
add_bullet(doc, 'משיכת רשימת הלקוחות התקועים; פילוח לפי שלב אימוץ (0–3)', rtl=True)
add_bullet(doc, 'שיחות עם 3 נציגי CS: אילו חסמים אתם שומעים הכי הרבה? מה ה-playbook הנוכחי?', rtl=True)
add_bullet(doc, 'זיהוי איזה פלח לקוחות מציג הכי הרבה נפח וסיכון הנחשה הגבוה ביותר', rtl=True)

add_heading(doc, 'שבוע 2 — אבחון החסמים האמיתיים', level=3, rtl=True)
add_bullet(doc, 'שיחות ישירות עם 5 לקוחות תקועים ב-2–3 פלחים', rtl=True)
add_bullet(doc, 'מטרה: מציאת 2–3 נקודות חיכוך שחוזרות על עצמן (למשל: "הגדרת מדיניות מורכבת מדי", '
           '"עומס CVE ללא תעדוף", "אין שיתוף פעולה של DevOps")', rtl=True)
add_bullet(doc, 'זיהוי 1–2 אנשי קשר "אלופים" בלקוחות תקועים שיסכימו לשתף פעולה בעיצוב פתרון', rtl=True)

add_heading(doc, 'שבוע 3 — הרצת ניסוי ממוקד אחד', level=3, rtl=True)
add_bullet(doc, 'הגדרת ההתערבות המינימלית לתבנית ה-stuck הנפוצה ביותר', rtl=True)
add_bullet(doc, 'הרצה עם 3 לקוחות: למשל, שיחת "Xray Activation Sprint" של 45 דקות '
           'עם תבנית מדיניות מוכנה מראש והדרכת הגדרה', rtl=True)
add_bullet(doc, 'מדידה: זמן עד ליצירת המדיניות הראשונה, זמן עד לטיפול ב-CVE הראשון', rtl=True)

add_heading(doc, 'שבוע 4 — הפצה ובניית ה-Playbook', level=3, rtl=True)
add_bullet(doc, 'הצגת תוצאות ראשוניות להנהלת CS', rtl=True)
add_bullet(doc, 'הצעת playbook מובנה אחד ל-CS: מה מפעיל פנייה, איך נראית השיחה, מה נחשב הצלחה', rtl=True)
add_bullet(doc, 'הגדרת ציון בריאות האימוץ שיזין את דשבורד ה-CS', rtl=True)

add_body(doc, 'המסגרת המרכזית: הסיכון הגדול ביותר אינו שלא ידעים מה לבנות — '
         'אלא של-CS אין תנועה חוזרת. 30 הימים הראשונים עוסקים ביצירת מספיק אות '
         'כדי לתת ל-CS playbook, לא בבניית מוצר מאפס.', rtl=True)

# =====================================================================
# SAVE
# =====================================================================
output_path = '/Users/assaf/My Stuff/Calude/Claude Code/Jfrog interview/Part 1.docx'
doc.save(output_path)
print(f"Saved: {output_path}")
